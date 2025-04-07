#!/usr/bin/env python3
"""
Script de démonstration pour DocxFilesMerger CLI
Ce script crée une archive ZIP de test, puis utilise le CLI pour la traiter.

Développé par MOA Digital Agency LLC (https://myoneart.com)
Email: moa@myoneart.com
Copyright © 2025 MOA Digital Agency LLC. Tous droits réservés.
"""

import os
import sys
import zipfile
import tempfile
import shutil
from docx import Document
from docx_merger_cli import traiter_fichier


def creer_docx_test(chemin_sortie, contenu="Ceci est un document de test."):
    """Crée un document DOCX de test simple"""
    doc = Document()
    doc.add_heading('Document de Test', 0)
    doc.add_paragraph(contenu)
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('Voici la première section du document test.')
    doc.add_heading('Section 2', level=1)
    doc.add_paragraph('Voici la deuxième section du document test.')
    
    # Sauvegarder le document
    doc.save(chemin_sortie)
    return chemin_sortie


def creer_archive_test(chemin_sortie, nb_fichiers=3):
    """
    Crée une archive ZIP contenant plusieurs documents DOCX pour les tests
    
    Args:
        chemin_sortie: Chemin où créer l'archive ZIP
        nb_fichiers: Nombre de fichiers DOCX à créer dans l'archive
        
    Returns:
        Le chemin vers l'archive ZIP créée
    """
    # Créer un dossier temporaire
    with tempfile.TemporaryDirectory() as temp_dir:
        # Créer plusieurs fichiers DOCX
        fichiers_docx = []
        for i in range(nb_fichiers):
            nom_fichier = f"document_test_{i+1}.docx"
            chemin_fichier = os.path.join(temp_dir, nom_fichier)
            contenu = f"Ceci est le document test numéro {i+1}.\n"
            contenu += f"Il contient du texte unique pour le tester.\n"
            contenu += f"Ce fichier sera fusionné avec les autres documents."
            creer_docx_test(chemin_fichier, contenu)
            fichiers_docx.append(chemin_fichier)
        
        # Créer l'archive ZIP
        with zipfile.ZipFile(chemin_sortie, 'w') as zipf:
            for fichier in fichiers_docx:
                zipf.write(fichier, os.path.basename(fichier))
                
        print(f"Archive ZIP créée: {chemin_sortie}")
        print(f"Contient {len(fichiers_docx)} fichiers DOCX")
        
    return chemin_sortie


def main():
    """Fonction principale de démonstration"""
    print("=== Démonstration du DocxFilesMerger CLI ===")
    
    # Créer un dossier pour les tests
    dossier_test = "./test_demo"
    os.makedirs(dossier_test, exist_ok=True)
    
    # Créer une archive ZIP de test
    archive_test = os.path.join(dossier_test, "archive_test.zip")
    creer_archive_test(archive_test, nb_fichiers=5)
    
    # Dossier pour les résultats
    dossier_resultats = os.path.join(dossier_test, "resultats")
    
    print("\n=== Traitement de l'archive avec DocxFilesMerger ===")
    
    # Traiter l'archive
    resultat = traiter_fichier(archive_test, dossier_resultats)
    
    # Afficher les résultats
    print("\n=== Résultats ===")
    print(f"Statut: {resultat['statut']}")
    print(f"Message: {resultat['message']}")
    print(f"Temps de traitement: {resultat.get('temps', 0):.2f} secondes")
    
    if resultat['docx']:
        print(f"Document DOCX créé: {resultat['docx']}")
    
    if resultat['pdf']:
        print(f"Document PDF créé: {resultat['pdf']}")
    
    print("\nLa démonstration est terminée.")
    
    return 0


if __name__ == "__main__":
    sys.exit(main())