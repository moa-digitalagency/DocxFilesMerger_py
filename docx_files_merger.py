#!/usr/bin/env python3
"""
DocxFilesMerger - Application de traitement et fusion de documents.
Développé par MOA Digital Agency LLC (https://myoneart.com)
Email: moa@myoneart.com
Copyright © 2025 MOA Digital Agency LLC. Tous droits réservés.
"""

import os
import sys
import zipfile
import tempfile
import shutil
import argparse
import subprocess
import time
from pathlib import Path

# Essai d'importation des dépendances optionnelles
try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False

# Importation des bibliothèques principales
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Erreur: La bibliothèque python-docx n'est pas installée.")
    print("Installez-la avec: pip install python-docx")
    sys.exit(1)

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    print("Avertissement: La bibliothèque reportlab n'est pas installée.")
    print("Le PDF généré sera de qualité basique.")
    print("Installez-la avec: pip install reportlab")


def print_progress(message, percent, step=""):
    """Affiche l'avancement du traitement dans le terminal"""
    bar_length = 40
    filled_length = int(bar_length * percent / 100)
    bar = '█' * filled_length + '░' * (bar_length - filled_length)
    
    # Afficher la barre de progression
    if step:
        status = f"[{step}]"
    else:
        status = ""
        
    sys.stdout.write(f"\r{message} {status} [{bar}] {percent:.1f}%")
    sys.stdout.flush()
    
    # Ajouter une nouvelle ligne si on a atteint 100%
    if percent >= 100:
        sys.stdout.write("\n")


def extract_doc_files(zip_path, extract_dir):
    """Extract all .doc and .docx files from a zip file"""
    # Ensure extract directory exists
    os.makedirs(extract_dir, exist_ok=True)
    
    # Open and extract files
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        all_files = zip_ref.namelist()
        doc_files = [f for f in all_files if f.lower().endswith(('.doc', '.docx'))]
        
        if not doc_files:
            print("Aucun fichier DOC/DOCX trouvé dans l'archive.")
            return []
        
        # Extract only .doc and .docx files
        for i, file_path in enumerate(doc_files):
            # Extract file
            zip_ref.extract(file_path, path=extract_dir)
            
            # Get just the filename without directory structure
            file_name = os.path.basename(file_path)
            
            # If the file was in a subdirectory, move it to the extract_dir
            extracted_path = os.path.join(extract_dir, file_path)
            if os.path.dirname(extracted_path) != extract_dir:
                target_path = os.path.join(extract_dir, file_name)
                os.makedirs(os.path.dirname(target_path), exist_ok=True)
                shutil.move(extracted_path, target_path)
                
                # Clean up empty directories
                dir_path = os.path.dirname(extracted_path)
                while dir_path != extract_dir:
                    try:
                        os.rmdir(dir_path)
                        dir_path = os.path.dirname(dir_path)
                    except OSError:
                        break

    # Get all extracted files
    extracted_files = []
    for root, _, files in os.walk(extract_dir):
        for file in files:
            if file.lower().endswith(('.doc', '.docx')):
                extracted_files.append(os.path.join(root, file))
    
    return extracted_files


def convert_doc_to_docx(doc_path, output_dir):
    """
    Convert a .doc file to .docx format
    
    This function attempts to use LibreOffice for conversion if available,
    falling back to a basic document creation method if not
    """
    # Skip if already a .docx file
    if doc_path.lower().endswith('.docx'):
        return doc_path
    
    # Prepare output path
    docx_filename = os.path.basename(doc_path).rsplit('.', 1)[0] + '.docx'
    docx_path = os.path.join(output_dir, docx_filename)
    
    # Try to use LibreOffice for conversion
    try:
        # Create a temporary directory
        with tempfile.TemporaryDirectory() as temp_dir:
            # Construct LibreOffice command
            command = [
                'libreoffice', '--headless', '--convert-to', 'docx',
                '--outdir', temp_dir, doc_path
            ]
            
            # Run LibreOffice conversion
            result = subprocess.run(command, capture_output=True, text=True)
            
            if result.returncode == 0:
                # Get the output file
                converted_file = os.path.join(temp_dir, docx_filename)
                
                # If conversion successful, move to the target location
                if os.path.exists(converted_file):
                    shutil.move(converted_file, docx_path)
                    return docx_path
    except (subprocess.SubprocessError, FileNotFoundError):
        # LibreOffice not available or conversion failed
        pass
    
    # Fallback: Create a new document with basic content
    print(f"\nAvertissement: Impossible de convertir {doc_path} avec LibreOffice.")
    print("Création d'un document DOCX basique à la place.")
    
    # Create a basic DOCX file
    doc = Document()
    doc.add_heading(f"Document converti: {os.path.basename(doc_path)}", 0)
    doc.add_paragraph("Ce document a été créé car la conversion automatique du fichier DOC original a échoué.")
    doc.add_paragraph("Pour une conversion complète, installez LibreOffice ou utilisez Microsoft Word.")
    doc.add_paragraph("")
    doc.add_paragraph("Informations sur le fichier d'origine:")
    doc.add_paragraph(f"- Nom: {os.path.basename(doc_path)}")
    doc.add_paragraph(f"- Chemin: {doc_path}")
    doc.add_paragraph(f"- Taille: {os.path.getsize(doc_path)} octets")
    
    # Save the document
    doc.save(docx_path)
    return docx_path


def merge_docx_files(docx_files, output_path):
    """
    Merge multiple .docx files into a single document
    
    Before each file's content, a header line with the filename is added.
    Updates progress in the terminal.
    """
    if not docx_files:
        print("Aucun fichier DOCX à fusionner.")
        return None

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Create a new document
    merged_doc = Document()
    merged_doc.add_heading('Documents Fusionnés', 0)
    
    # Add table of contents
    merged_doc.add_heading('Table des matières', level=1)
    for i, doc_path in enumerate(docx_files, 1):
        filename = os.path.basename(doc_path)
        toc_para = merged_doc.add_paragraph(f"{i}. ")
        toc_para.add_run(filename).bold = True
    
    merged_doc.add_page_break()
    
    # Define progress tracking
    total_files = len(docx_files)
    
    # Process each document
    for i, doc_path in enumerate(docx_files):
        # Update progress
        percent = (i / total_files) * 100
        print_progress("Fusion des documents", percent, f"{i+1}/{total_files}")
        
        try:
            # Document header
            filename = os.path.basename(doc_path)
            merged_doc.add_heading(f"Document {i+1}: {filename}", level=1)
            
            # Open the document to merge
            src_doc = Document(doc_path)
            
            # Copy all elements from the source document
            for element in src_doc.element.body:
                merged_doc.element.body.append(element)
            
            # Add a page break after each document except the last one
            if i < len(docx_files) - 1:
                merged_doc.add_page_break()
                
        except Exception as e:
            print(f"\nErreur lors de la fusion du document {doc_path}: {str(e)}")
            # Continue with the next document
    
    # Final progress update
    print_progress("Fusion des documents", 100, f"{total_files}/{total_files}")
    
    # Save the merged document
    try:
        merged_doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"\nErreur lors de l'enregistrement du document fusionné: {str(e)}")
        return None


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Convert a .docx file to .pdf format
    
    This function attempts multiple methods to convert the document:
    1. libreoffice (if available)
    2. docx2pdf library (if installed)
    3. Basic fallback message if conversion is not possible
    """
    # Ensure the output directory exists
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    
    # Try Method 1: LibreOffice
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            command = [
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', temp_dir, docx_path
            ]
            
            # Run the conversion
            print("Conversion PDF via LibreOffice...")
            result = subprocess.run(command, capture_output=True, text=True)
            
            if result.returncode == 0:
                # Get the output file name
                pdf_filename = os.path.basename(docx_path).rsplit('.', 1)[0] + '.pdf'
                temp_pdf = os.path.join(temp_dir, pdf_filename)
                
                # Move to the target location if successful
                if os.path.exists(temp_pdf):
                    shutil.move(temp_pdf, pdf_path)
                    return pdf_path
    except (subprocess.SubprocessError, FileNotFoundError):
        # LibreOffice not available
        pass
    
    # Try Method 2: docx2pdf library
    try:
        import docx2pdf
        print("Conversion PDF via docx2pdf...")
        docx2pdf.convert(docx_path, pdf_path)
        
        if os.path.exists(pdf_path):
            return pdf_path
    except ImportError:
        # docx2pdf not available
        pass
    except Exception as e:
        print(f"Erreur lors de la conversion avec docx2pdf: {str(e)}")
    
    # Method 3: Basic PDF with reportlab
    try:
        # Open the Word document
        doc = Document(docx_path)
        
        # Create a PDF
        c = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter
        
        # Extract and write text
        c.setFont("Helvetica", 16)
        c.drawString(72, height - 72, "Document converti en PDF")
        
        c.setFont("Helvetica", 12)
        c.drawString(72, height - 100, "Ce PDF contient uniquement le texte extrait du document DOCX.")
        c.drawString(72, height - 120, "Pour un PDF complet avec mise en forme, installez LibreOffice.")
        
        y_position = height - 160
        c.setFont("Helvetica", 10)
        
        # Add paragraphs
        for para in doc.paragraphs:
            if para.text:
                # Wrap text to fit page width
                text = para.text
                while text and y_position > 72:
                    if len(text) > 80:  # Simple text wrapping
                        wrapped_text = text[:80]
                        text = text[80:]
                    else:
                        wrapped_text = text
                        text = ""
                    
                    c.drawString(72, y_position, wrapped_text)
                    y_position -= 15
                
                # Add extra space between paragraphs
                y_position -= 5
            
            # Start a new page if needed
            if y_position < 72:
                c.showPage()
                y_position = height - 72
                c.setFont("Helvetica", 10)
        
        # Save the PDF
        c.save()
        print("PDF basique créé avec ReportLab.")
        return pdf_path
        
    except Exception as e:
        print(f"Erreur lors de la création du PDF basique: {str(e)}")
    
    # If all methods failed
    print("Impossible de convertir le document en PDF.")
    return None


def process_zip_file(zip_path, output_dir, show_progress=True):
    """
    Process a zip file containing .doc/.docx files:
    1. Extract all .doc and .docx files
    2. Convert .doc to .docx if needed
    3. Merge all into a single .docx
    4. Convert the merged file to PDF
    
    Returns a tuple of (docx_path, pdf_path) with the paths to the generated files
    """
    # Create output directories
    os.makedirs(output_dir, exist_ok=True)
    extract_dir = os.path.join(output_dir, "extracted")
    os.makedirs(extract_dir, exist_ok=True)
    
    # Define output paths
    docx_output = os.path.join(output_dir, "merged.docx")
    pdf_output = os.path.join(output_dir, "merged.pdf")
    
    # Process files
    if show_progress:
        print(f"Traitement du fichier ZIP: {os.path.basename(zip_path)}")
    
    # Step 1: Extract files
    if show_progress:
        print("Étape 1: Extraction des fichiers...")
    extracted_files = extract_doc_files(zip_path, extract_dir)
    
    if not extracted_files:
        print("Aucun fichier DOC/DOCX trouvé dans l'archive.")
        return None, None
    
    if show_progress:
        print(f"  ✓ {len(extracted_files)} fichiers extraits")
    
    # Step 2: Convert .doc to .docx if needed
    if show_progress:
        print("Étape 2: Conversion des fichiers DOC en DOCX...")
    
    docx_files = []
    doc_files = [f for f in extracted_files if f.lower().endswith('.doc')]
    
    if doc_files:
        if show_progress:
            print(f"  Conversion de {len(doc_files)} fichiers DOC...")
        
        for i, doc_path in enumerate(doc_files):
            if show_progress:
                percent = ((i + 1) / len(doc_files)) * 100
                print_progress("Conversion DOC vers DOCX", percent, f"{i+1}/{len(doc_files)}")
            
            docx_path = convert_doc_to_docx(doc_path, extract_dir)
            docx_files.append(docx_path)
    
    # Add already .docx files
    docx_files.extend([f for f in extracted_files if f.lower().endswith('.docx')])
    
    # Sort files by name for consistent ordering
    docx_files.sort()
    
    if show_progress:
        print(f"  ✓ {len(docx_files)} fichiers DOCX prêts pour la fusion")
    
    # Step 3: Merge all .docx files
    if show_progress:
        print("Étape 3: Fusion des documents...")
    
    merged_docx = merge_docx_files(docx_files, docx_output)
    
    if not merged_docx:
        print("Erreur lors de la fusion des documents.")
        return None, None
    
    if show_progress:
        print(f"  ✓ Document fusionné créé: {os.path.basename(merged_docx)}")
    
    # Step 4: Convert to PDF
    if show_progress:
        print("Étape 4: Conversion en PDF...")
    
    pdf_path = convert_docx_to_pdf(merged_docx, pdf_output)
    
    if pdf_path:
        if show_progress:
            print(f"  ✓ PDF créé: {os.path.basename(pdf_path)}")
    else:
        if show_progress:
            print("  ✗ Échec de la conversion en PDF")
    
    # Summary
    if show_progress:
        print("\nTraitement terminé!")
        print(f"Documents traités: {len(docx_files)}")
        print(f"Document DOCX fusionné: {merged_docx}")
        if pdf_path:
            print(f"Document PDF: {pdf_path}")
    
    return merged_docx, pdf_path


def main():
    """Fonction principale pour le traitement des fichiers via ligne de commande"""
    # Parser les arguments de ligne de commande
    parser = argparse.ArgumentParser(
        description="Fusionne des fichiers DOC/DOCX depuis une archive ZIP en un seul document.",
        epilog="Développé par MOA Digital Agency LLC (https://myoneart.com)"
    )
    parser.add_argument("zip_file", help="Chemin vers le fichier ZIP contenant les documents")
    parser.add_argument("-o", "--output-dir", default="./output", 
                        help="Dossier de sortie pour les fichiers générés (par défaut: ./output)")
    parser.add_argument("-q", "--quiet", action="store_true", 
                        help="Mode silencieux (pas d'affichage des barres de progression)")
    
    args = parser.parse_args()
    
    # Vérifier que le fichier ZIP existe
    if not os.path.isfile(args.zip_file):
        print(f"Erreur: Le fichier {args.zip_file} n'existe pas.")
        return 1
    
    # Traiter le fichier
    start_time = time.time()
    
    try:
        docx_path, pdf_path = process_zip_file(
            args.zip_file, 
            args.output_dir,
            show_progress=not args.quiet
        )
        
        processing_time = time.time() - start_time
        
        if not args.quiet:
            print(f"\nTemps de traitement: {processing_time:.2f} secondes")
        
        if docx_path and pdf_path:
            return 0
        else:
            return 1
    
    except Exception as e:
        print(f"Erreur: {str(e)}")
        return 1


if __name__ == "__main__":
    sys.exit(main())