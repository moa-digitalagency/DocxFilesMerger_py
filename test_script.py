#!/usr/bin/env python3
"""
Test simple du script docx_files_merger.py

Ce script crée un fichier ZIP avec un document DOCX de test,
puis utilise le script docx_files_merger.py pour le traiter.

Développé par MOA Digital Agency LLC (https://myoneart.com)
Email: moa@myoneart.com
Copyright © 2025 MOA Digital Agency LLC. Tous droits réservés.
"""

import os
import sys
import zipfile
import tempfile
import shutil
from pathlib import Path
try:
    from docx import Document
except ImportError:
    print("Erreur: La bibliothèque python-docx n'est pas installée.")
    print("Installez-la avec: pip install python-docx")
    sys.exit(1)

# Importer le module de traitement
from docx_files_merger import process_zip_file


def create_test_docx(output_path):
    """Crée un simple document DOCX de test"""
    doc = Document()
    
    # Ajouter un titre
    doc.add_heading('Document de Test', 0)
    
    # Ajouter du texte en paragraphes
    doc.add_paragraph('Ceci est un document de test généré automatiquement.')
    
    p = doc.add_paragraph('Ce document est utilisé pour tester ')
    p.add_run('DocxFilesMerger CLI').bold = True
    p.add_run(', un outil de fusion de documents DOCX.')
    
    # Ajouter une liste à puces
    doc.add_heading('Fonctionnalités testées:', level=1)
    doc.add_paragraph('Extraction depuis une archive ZIP', style='List Bullet')
    doc.add_paragraph('Fusion de plusieurs documents DOCX', style='List Bullet')
    doc.add_paragraph('Génération de PDF à partir du document fusionné', style='List Bullet')
    
    # Créer le dossier de sortie si nécessaire
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    
    # Sauvegarder le document
    doc.save(output_path)
    return output_path


def create_test_zip(output_path, num_files=3):
    """
    Crée une archive ZIP contenant plusieurs documents DOCX de test
    
    Args:
        output_path: Chemin où créer l'archive ZIP
        num_files: Nombre de fichiers à créer dans l'archive
    
    Returns:
        Le chemin vers l'archive ZIP créée
    """
    # Créer un dossier temporaire
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Créer plusieurs documents DOCX
        docx_files = []
        for i in range(1, num_files + 1):
            docx_path = os.path.join(temp_dir, f"document_test_{i}.docx")
            create_test_docx(docx_path)
            docx_files.append(docx_path)
        
        # Créer le dossier de sortie si nécessaire
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        
        # Créer l'archive ZIP
        with zipfile.ZipFile(output_path, 'w') as zipf:
            for docx_file in docx_files:
                zipf.write(docx_file, os.path.basename(docx_file))
        
        print(f"Archive ZIP créée avec {num_files} fichiers DOCX: {output_path}")
        return output_path
        
    finally:
        # Nettoyer le dossier temporaire
        shutil.rmtree(temp_dir)


def main():
    """Fonction principale de test"""
    # Définir les chemins
    test_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "test_docx_merger")
    zip_path = os.path.join(test_dir, "archive_test.zip")
    output_dir = os.path.join(test_dir, "resultats")
    
    # Créer les dossiers si nécessaire
    os.makedirs(test_dir, exist_ok=True)
    
    # Créer une archive ZIP de test
    create_test_zip(zip_path, num_files=5)
    
    # Traiter l'archive
    print("\nTraitement de l'archive de test...")
    docx_path, pdf_path = process_zip_file(zip_path, output_dir)
    
    # Vérifier les résultats
    if docx_path and pdf_path and os.path.exists(docx_path) and os.path.exists(pdf_path):
        print("\nTest réussi!")
        print(f"Document DOCX fusionné: {docx_path}")
        print(f"Document PDF généré: {pdf_path}")
        print(f"\nFichiers du test disponibles dans: {test_dir}")
        return 0
    else:
        print("\nTest échoué!")
        print("Les fichiers n'ont pas été générés correctement.")
        return 1


if __name__ == "__main__":
    sys.exit(main())