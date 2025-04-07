import os
import shutil
import zipfile
import json
import time
import threading
import traceback
from datetime import datetime, timedelta
from pathlib import Path
import subprocess
import sys
import tempfile

# Import des bibliothèques de traitement de documents
try:
    import docx
    from docx import Document
except ImportError:
    print("Bibliothèque python-docx non installée. Certaines fonctionnalités peuvent ne pas fonctionner correctement.")

def save_status(status_dir, status_data):
    """Save processing status to a JSON file"""
    if not status_dir:
        return

    os.makedirs(status_dir, exist_ok=True)
    status_file = os.path.join(status_dir, 'status.json')
    
    try:
        with open(status_file, 'w') as f:
            json.dump(status_data, f)
    except Exception as e:
        print(f"Erreur lors de la sauvegarde du statut: {str(e)}")

def extract_doc_files(zip_path, extract_dir):
    """Extract all .doc and .docx files from a zip file"""
    # Créer le dossier d'extraction s'il n'existe pas
    os.makedirs(extract_dir, exist_ok=True)
    
    # Liste pour stocker les chemins vers les fichiers extraits
    extracted_files = []
    
    # Ouvrir le fichier ZIP
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        # Parcourir tous les fichiers dans le ZIP
        for file_info in zip_ref.infolist():
            # Ignorer les dossiers
            if file_info.filename.endswith('/'):
                continue
            
            # Vérifier si le fichier est un .doc ou .docx
            if file_info.filename.lower().endswith(('.doc', '.docx')):
                # Extraire uniquement le nom du fichier sans les dossiers
                filename = os.path.basename(file_info.filename)
                # Créer un chemin de destination
                dest_path = os.path.join(extract_dir, filename)
                
                # Extraire le fichier
                with zip_ref.open(file_info) as source, open(dest_path, 'wb') as dest:
                    shutil.copyfileobj(source, dest)
                
                # Ajouter le chemin à la liste
                extracted_files.append(dest_path)
    
    return extracted_files

def convert_doc_to_docx(doc_path, output_dir):
    """
    Convert a .doc file to .docx format
    
    This function attempts to use LibreOffice for conversion if available,
    falling back to a basic document creation method if not
    """
    try:
        # Créer une instance de Document pour vérifier le chargement
        document = Document(doc_path)
        
        # Si ça fonctionne, simplement sauvegarder en .docx
        filename = os.path.basename(doc_path)
        name_without_ext = os.path.splitext(filename)[0]
        docx_path = os.path.join(output_dir, f"{name_without_ext}.docx")
        
        document.save(docx_path)
        return docx_path
        
    except Exception as e:
        print(f"Échec de la conversion directe de {doc_path}: {str(e)}")
        
        try:
            # Essayer de convertir avec LibreOffice si disponible
            # LibreOffice doit être installé sur le système
            filename = os.path.basename(doc_path)
            name_without_ext = os.path.splitext(filename)[0]
            docx_path = os.path.join(output_dir, f"{name_without_ext}.docx")
            
            cmd = ['libreoffice', '--headless', '--convert-to', 'docx', 
                   '--outdir', output_dir, doc_path]
            
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            
            # Vérifier si le fichier a bien été créé
            if os.path.exists(docx_path):
                return docx_path
            
        except (subprocess.SubprocessError, FileNotFoundError) as e:
            print(f"Échec de la conversion via LibreOffice: {str(e)}")
            
            try:
                # Si tout échoue, créer un nouveau document avec le contenu
                import docx2pdf
                print("Tentative de conversion avec docx2pdf...")
                
                # Obtenir le chemin de sortie
                filename = os.path.basename(doc_path)
                name_without_ext = os.path.splitext(filename)[0]
                docx_path = os.path.join(output_dir, f"{name_without_ext}.docx")
                
                # Créer un document vierge avec un message
                doc = Document()
                doc.add_paragraph(f"Le fichier {filename} n'a pas pu être converti automatiquement.")
                doc.add_paragraph("Veuillez consulter le fichier original.")
                doc.save(docx_path)
                
                return docx_path
                
            except Exception as inner_e:
                print(f"Échec de la création d'un document de remplacement: {str(inner_e)}")
                return None

def merge_docx_files(docx_files, output_path, status_dir):
    """
    Merge multiple .docx files into a single document
    
    Before each file's content, a header line with the filename is added.
    Updates status periodically.
    """
    # S'assurer que nous avons des fichiers à fusionner
    if not docx_files:
        save_status(status_dir, {
            'percent': 100,
            'status_text': 'Aucun fichier à fusionner trouvé.',
            'current_step': 'merge',
            'complete': True,
            'error': 'Aucun fichier DOCX trouvé'
        })
        return None
    
    # Créer un nouveau document
    merged_doc = Document()
    
    total_files = len(docx_files)
    
    # Parcourir chaque fichier
    for index, file_path in enumerate(docx_files):
        # Mettre à jour le statut
        progress = int((index / total_files) * 100)
        save_status(status_dir, {
            'percent': progress,
            'status_text': f'Fusion du document {index+1}/{total_files}...',
            'current_step': 'merge',
            'complete': False,
            'file_count': total_files
        })
        
        try:
            # Obtenir le nom du fichier
            filename = os.path.basename(file_path)
            
            # Ajouter un saut de page si ce n'est pas le premier document
            if index > 0:
                merged_doc.add_page_break()
            
            # Ajouter une section d'en-tête avec le nom du fichier
            merged_doc.add_heading(f'Document: {filename}', level=1)
            
            # Ouvrir le document source
            src_doc = Document(file_path)
            
            # Copier tous les paragraphes
            for paragraph in src_doc.paragraphs:
                p = merged_doc.add_paragraph()
                for run in paragraph.runs:
                    p.add_run(run.text, run.style)
            
            # Pour les éléments plus complexes, un traitement spécifique serait nécessaire
            # (tableaux, images, etc.)
            
        except Exception as e:
            print(f"Erreur lors de la fusion du fichier {file_path}: {str(e)}")
            # Ajouter un paragraphe d'erreur
            merged_doc.add_paragraph(f"Erreur lors de la fusion du fichier {filename}: {str(e)}")
    
    # Sauvegarder le document fusionné
    try:
        merged_doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"Erreur lors de la sauvegarde du document fusionné: {str(e)}")
        save_status(status_dir, {
            'percent': 0,
            'status_text': 'Erreur lors de la sauvegarde du document fusionné.',
            'current_step': 'error',
            'complete': False,
            'error': str(e)
        })
        return None

def convert_docx_to_pdf(docx_path, pdf_path, status_dir):
    """
    Convert a .docx file to .pdf format
    
    This function attempts multiple methods to convert the document:
    1. libreoffice (if available)
    2. docx2pdf library (if installed)
    3. Basic fallback message if conversion is not possible
    """
    if not os.path.exists(docx_path):
        save_status(status_dir, {
            'percent': 0,
            'status_text': 'Le fichier DOCX à convertir n\'existe pas.',
            'current_step': 'error',
            'complete': False,
            'error': 'Fichier DOCX introuvable'
        })
        return None
    
    save_status(status_dir, {
        'percent': 70,
        'status_text': 'Conversion en PDF...',
        'current_step': 'pdf',
        'complete': False
    })
    
    # Méthode 1: Essayer avec LibreOffice
    try:
        cmd = ['libreoffice', '--headless', '--convert-to', 'pdf', 
               '--outdir', os.path.dirname(pdf_path), docx_path]
        
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        # LibreOffice génère le PDF avec le même nom de base que l'entrée
        generated_pdf = os.path.splitext(docx_path)[0] + '.pdf'
        
        # Si le chemin de sortie est différent, renommer le fichier
        if generated_pdf != pdf_path and os.path.exists(generated_pdf):
            shutil.move(generated_pdf, pdf_path)
        
        if os.path.exists(pdf_path):
            return pdf_path
        
    except (subprocess.SubprocessError, FileNotFoundError) as e:
        print(f"Échec de la conversion PDF via LibreOffice: {str(e)}")
    
    # Méthode 2: Essayer avec docx2pdf
    try:
        import docx2pdf
        docx2pdf.convert(docx_path, pdf_path)
        
        if os.path.exists(pdf_path):
            return pdf_path
            
    except ImportError:
        print("Bibliothèque docx2pdf non installée.")
    except Exception as e:
        print(f"Échec de la conversion PDF via docx2pdf: {str(e)}")
    
    # Méthode 3: Créer un PDF basique avec reportlab
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.setFont("Helvetica", 12)
        
        # Message d'information
        c.drawString(100, 700, "Conversion automatique du document DOCX en PDF")
        c.drawString(100, 680, "La conversion automatique n'a pas pu être effectuée.")
        c.drawString(100, 660, "Veuillez utiliser le fichier DOCX fourni.")
        
        c.save()
        
        if os.path.exists(pdf_path):
            return pdf_path
            
    except ImportError:
        print("Bibliothèque reportlab non installée.")
    except Exception as e:
        print(f"Échec de la création d'un PDF basique: {str(e)}")
    
    # Si tout échoue, créer un fichier texte simple
    try:
        txt_path = os.path.splitext(pdf_path)[0] + '.txt'
        with open(txt_path, 'w') as f:
            f.write("ERREUR DE CONVERSION PDF\n\n")
            f.write("La conversion automatique du document DOCX en PDF n'a pas pu être effectuée.\n")
            f.write("Veuillez utiliser le fichier DOCX fourni ou installer les outils nécessaires (LibreOffice, docx2pdf ou reportlab).")
        
        # Essayer une dernière fois avec reportlab
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            
            c = canvas.Canvas(pdf_path, pagesize=letter)
            c.setFont("Helvetica", 12)
            
            with open(txt_path, 'r') as f:
                y = 700
                for line in f:
                    c.drawString(100, y, line.strip())
                    y -= 20
            
            c.save()
            
            if os.path.exists(pdf_path):
                return pdf_path
                
        except Exception:
            pass
            
        return txt_path
        
    except Exception as e:
        print(f"Échec de la création d'un fichier texte: {str(e)}")
    
    return None

def process_zip_file(zip_path, output_dir, status_dir=None, job_id=None):
    """
    Process a zip file containing .doc/.docx files:
    1. Extract all .doc and .docx files
    2. Convert .doc to .docx if needed
    3. Merge all into a single .docx
    4. Convert the merged file to PDF
    
    This function operates asynchronously and updates a status file.
    If job_id is provided, it will update the database with processing status.
    """
    # Function to be run in a separate thread
    def process_thread():
        start_time = int(time.time())
        
        try:
            # Créer les dossiers de sortie
            os.makedirs(output_dir, exist_ok=True)
            if status_dir:
                os.makedirs(status_dir, exist_ok=True)
            
            # Étape 1: Extraction des fichiers
            save_status(status_dir, {
                'percent': 10,
                'status_text': 'Extraction des fichiers...',
                'current_step': 'extract',
                'complete': False,
                'start_time': start_time
            })
            
            # Extraire les fichiers .doc et .docx
            extract_folder = os.path.join(output_dir, 'extracted')
            extracted_files = extract_doc_files(zip_path, extract_folder)
            
            if not extracted_files:
                save_status(status_dir, {
                    'percent': 0,
                    'status_text': 'Aucun fichier DOC/DOCX trouvé dans l\'archive ZIP.',
                    'current_step': 'error',
                    'complete': False,
                    'error': 'Archive vide ou sans fichiers DOC/DOCX',
                    'start_time': start_time,
                    'end_time': int(time.time())
                })
                
                # Mettre à jour le statut dans la base de données
                if job_id:
                    try:
                        import sys
                        sys.path.append(os.getcwd())
                        from flask import Flask
                        from models import db, ProcessingJob
                        
                        app = Flask(__name__)
                        app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL')
                        app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
                        db.init_app(app)
                        
                        with app.app_context():
                            job = ProcessingJob.query.filter_by(job_id=job_id).first()
                            if job:
                                job.status = 'error'
                                job.completed_at = datetime.now()
                                db.session.commit()
                    except Exception as db_err:
                        print(f"Erreur lors de la mise à jour du statut dans la base de données: {str(db_err)}")
                
                return
            
            # Étape 2: Conversion des fichiers .doc en .docx
            save_status(status_dir, {
                'percent': 30,
                'status_text': 'Conversion des fichiers DOC en DOCX...',
                'current_step': 'convert',
                'complete': False,
                'start_time': start_time
            })
            
            # Liste pour les fichiers DOCX (convertis ou originaux)
            docx_files = []
            
            for file_path in extracted_files:
                if file_path.lower().endswith('.doc'):
                    # Convertir en DOCX
                    docx_path = convert_doc_to_docx(file_path, extract_folder)
                    if docx_path:
                        docx_files.append(docx_path)
                elif file_path.lower().endswith('.docx'):
                    # Déjà au format DOCX
                    docx_files.append(file_path)
            
            # Étape 3: Fusion des fichiers DOCX
            save_status(status_dir, {
                'percent': 50,
                'status_text': 'Fusion des documents...',
                'current_step': 'merge',
                'complete': False,
                'file_count': len(docx_files),
                'start_time': start_time
            })
            
            # Fusionner les fichiers DOCX
            merged_docx_path = os.path.join(output_dir, 'merged.docx')
            merge_result = merge_docx_files(docx_files, merged_docx_path, status_dir)
            
            if not merge_result:
                save_status(status_dir, {
                    'percent': 0,
                    'status_text': 'Échec de la fusion des documents.',
                    'current_step': 'error',
                    'complete': False,
                    'error': 'Erreur lors de la fusion des fichiers DOCX',
                    'start_time': start_time,
                    'end_time': int(time.time())
                })
                
                # Mettre à jour le statut dans la base de données
                if job_id:
                    try:
                        import sys
                        sys.path.append(os.getcwd())
                        from flask import Flask
                        from models import db, ProcessingJob
                        
                        app = Flask(__name__)
                        app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL')
                        app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
                        db.init_app(app)
                        
                        with app.app_context():
                            job = ProcessingJob.query.filter_by(job_id=job_id).first()
                            if job:
                                job.status = 'error'
                                job.completed_at = datetime.now()
                                db.session.commit()
                    except Exception as db_err:
                        print(f"Erreur lors de la mise à jour du statut dans la base de données: {str(db_err)}")
                    
                return
            
            # Étape 4: Conversion en PDF
            save_status(status_dir, {
                'percent': 70,
                'status_text': 'Conversion en PDF...',
                'current_step': 'pdf',
                'complete': False,
                'start_time': start_time
            })
            
            # Convertir en PDF
            pdf_path = os.path.join(output_dir, 'merged.pdf')
            pdf_result = convert_docx_to_pdf(merged_docx_path, pdf_path, status_dir)
            
            # Terminer
            end_time = int(time.time())
            processing_time = end_time - start_time
            
            save_status(status_dir, {
                'percent': 100,
                'status_text': 'Traitement terminé avec succès.',
                'current_step': 'complete',
                'complete': True,
                'file_count': len(docx_files),
                'output_docx': os.path.basename(merged_docx_path),
                'output_pdf': os.path.basename(pdf_result) if pdf_result else None,
                'start_time': start_time,
                'end_time': end_time,
                'processing_time': processing_time
            })
            
            # Mettre à jour le statut dans la base de données
            if job_id:
                try:
                    import sys
                    sys.path.append(os.getcwd())
                    from flask import Flask
                    from models import db, ProcessingJob, UsageStat
                    
                    app = Flask(__name__)
                    app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL')
                    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
                    db.init_app(app)
                    
                    with app.app_context():
                        # Mettre à jour le job
                        job = ProcessingJob.query.filter_by(job_id=job_id).first()
                        if job:
                            job.status = 'completed'
                            job.completed_at = datetime.now()
                            job.file_count = len(docx_files)
                            job.processing_time = processing_time
                            db.session.commit()
                        
                        # Mettre à jour les statistiques d'utilisation
                        today = datetime.now().date()
                        usage_stat = UsageStat.query.filter_by(date=today).first()
                        
                        if usage_stat:
                            usage_stat.total_jobs += 1
                            usage_stat.total_files_processed += len(docx_files)
                            usage_stat.total_processing_time += processing_time
                        else:
                            usage_stat = UsageStat(
                                date=today,
                                total_jobs=1,
                                total_files_processed=len(docx_files),
                                total_processing_time=processing_time
                            )
                            db.session.add(usage_stat)
                        
                        db.session.commit()
                except Exception as db_err:
                    print(f"Erreur lors de la mise à jour du statut dans la base de données: {str(db_err)}")
            
        except Exception as e:
            error_message = str(e)
            error_traceback = traceback.format_exc()
            print(f"Error processing ZIP file: {error_message}")
            print(error_traceback)
            
            error_time = int(time.time())
            
            save_status(status_dir, {
                'percent': 0,
                'status_text': 'Une erreur s\'est produite.',
                'current_step': 'error',
                'complete': False,
                'error': error_message,
                'traceback': error_traceback,
                'start_time': start_time,
                'end_time': error_time
            })
            
            # Mettre à jour le job dans la base de données en cas d'erreur
            if job_id:
                try:
                    import sys
                    sys.path.append(os.getcwd())
                    from flask import Flask
                    from models import db, ProcessingJob
                    
                    app = Flask(__name__)
                    app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL')
                    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
                    db.init_app(app)
                    
                    with app.app_context():
                        job = ProcessingJob.query.filter_by(job_id=job_id).first()
                        if job:
                            job.status = 'error'
                            job.completed_at = datetime.now()
                            db.session.commit()
                except Exception as db_err:
                    print(f"Erreur lors de la mise à jour du statut d'erreur dans la base de données: {str(db_err)}")
    
    # Start processing in a thread
    thread = threading.Thread(target=process_thread)
    thread.daemon = True
    thread.start()
    
    return thread

def cleanup_old_files(directory, max_age_hours=24):
    """Delete files older than max_age_hours from the directory"""
    current_time = datetime.now()
    
    try:
        for item in os.listdir(directory):
            item_path = os.path.join(directory, item)
            
            # Skip if it's not a directory or file
            if not os.path.isdir(item_path) and not os.path.isfile(item_path):
                continue
                
            # Get the modification time
            mod_time = datetime.fromtimestamp(os.path.getmtime(item_path))
            
            # Check if the file/directory is older than max_age_hours
            if current_time - mod_time > timedelta(hours=max_age_hours):
                if os.path.isdir(item_path):
                    shutil.rmtree(item_path, ignore_errors=True)
                else:
                    os.remove(item_path)
                    
    except Exception as e:
        print(f"Error during cleanup: {str(e)}")
